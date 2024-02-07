import docx
import requests

country_document = docx.Document()  # creating the Word document - in memory of program

countries_response = requests.get('https://country-list-1150.herokuapp.com/api/country').json()  # call api

# Add a title to word document
country_document.add_paragraph('Countries and their Capital Cities ', 'Title')

# for loop adds selected data from api into word doc as long as data exists
for country_info in countries_response:
    print(country_info)
    name = country_info['name']  # access country name from api
    capital_city = country_info['capitalCity']  # access capital city name from api

    country_document.add_paragraph(f'The capital of {name} is {capital_city}.')  # format paragraph to create
    # sentence using name and capital city from api

# ask document to be saved and created named file on computer
country_document.save('countries.docx')
